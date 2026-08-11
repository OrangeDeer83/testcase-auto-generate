import io

from docx import Document


def extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)
